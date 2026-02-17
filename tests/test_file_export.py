"""Tests for the file export utilities (save_as_docx, save_as_text)."""

import os
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

from src.utils.file_export import save_as_docx, save_as_text


# ===================================================================
# save_as_text
# ===================================================================

class TestSaveAsText:

    def test_creates_file(self, tmp_path):
        out = tmp_path / "output.txt"
        result = save_as_text("hello world", str(out))
        assert result.exists()

    def test_returns_path_object(self, tmp_path):
        out = tmp_path / "output.txt"
        result = save_as_text("hello", str(out))
        assert isinstance(result, Path)

    def test_file_content_matches(self, tmp_path):
        out = tmp_path / "output.txt"
        save_as_text("hello world", str(out))
        assert out.read_text(encoding="utf-8") == "hello world"

    def test_multiline_content(self, tmp_path):
        content = "line1\nline2\nline3"
        out = tmp_path / "output.txt"
        save_as_text(content, str(out))
        assert out.read_text(encoding="utf-8") == content

    def test_empty_content(self, tmp_path):
        out = tmp_path / "output.txt"
        save_as_text("", str(out))
        assert out.read_text(encoding="utf-8") == ""

    def test_unicode_content(self, tmp_path):
        content = "Resume for Jose Garcia"
        out = tmp_path / "output.txt"
        save_as_text(content, str(out))
        assert out.read_text(encoding="utf-8") == content

    def test_creates_parent_directories(self, tmp_path):
        out = tmp_path / "nested" / "deep" / "output.txt"
        save_as_text("content", str(out))
        assert out.exists()

    def test_overwrites_existing_file(self, tmp_path):
        out = tmp_path / "output.txt"
        save_as_text("first", str(out))
        save_as_text("second", str(out))
        assert out.read_text(encoding="utf-8") == "second"

    def test_returned_path_matches_output_path(self, tmp_path):
        out = tmp_path / "output.txt"
        result = save_as_text("hello", str(out))
        assert result == out


# ===================================================================
# save_as_docx
# ===================================================================

class TestSaveAsDocx:

    def test_creates_file(self, tmp_path):
        out = tmp_path / "resume.docx"
        result = save_as_docx("Hello world", str(out))
        assert result.exists()

    def test_returns_path_object(self, tmp_path):
        out = tmp_path / "resume.docx"
        result = save_as_docx("Hello", str(out))
        assert isinstance(result, Path)

    def test_returned_path_matches_output_path(self, tmp_path):
        out = tmp_path / "resume.docx"
        result = save_as_docx("Hello", str(out))
        assert result == out

    def test_creates_parent_directories(self, tmp_path):
        out = tmp_path / "nested" / "dir" / "resume.docx"
        save_as_docx("content", str(out))
        assert out.exists()

    def test_file_is_valid_docx(self, tmp_path):
        from docx import Document

        out = tmp_path / "resume.docx"
        save_as_docx("Hello paragraph", str(out))
        doc = Document(str(out))
        assert len(doc.paragraphs) >= 1

    def test_single_paragraph(self, tmp_path):
        from docx import Document

        out = tmp_path / "resume.docx"
        save_as_docx("Single paragraph", str(out))
        doc = Document(str(out))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "Single paragraph" in texts

    def test_multiple_paragraphs(self, tmp_path):
        from docx import Document

        content = "Paragraph one\n\nParagraph two\n\nParagraph three"
        out = tmp_path / "resume.docx"
        save_as_docx(content, str(out))
        doc = Document(str(out))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert len(texts) == 3

    def test_empty_content(self, tmp_path):
        out = tmp_path / "empty.docx"
        result = save_as_docx("", str(out))
        assert result.exists()

    def test_overwrites_existing_file(self, tmp_path):
        from docx import Document

        out = tmp_path / "resume.docx"
        save_as_docx("First version", str(out))
        save_as_docx("Second version", str(out))
        doc = Document(str(out))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "Second version" in texts
        assert "First version" not in texts

    def test_file_size_is_nonzero(self, tmp_path):
        out = tmp_path / "resume.docx"
        save_as_docx("Some content", str(out))
        assert out.stat().st_size > 0
